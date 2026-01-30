import io
import os
import logging

import streamlit as st
from groq import Groq
import google.generativeai as genai
from PIL import Image as PILImage

# ------------------ ENV SETUP ------------------
from dotenv import load_dotenv
load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY", st.secrets.get("GROQ_API_KEY"))
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", st.secrets.get("GEMINI_API_KEY"))

# ------------------ RAG IMPORTS ------------------
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

# ------------------ PDF PARSING ------------------
try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader

# ------------------ DOCX EXPORT ------------------
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------ LOGGING ------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# =========================================================
# PDF HELPERS
# =========================================================

def extract_text_from_pdf(file_like):
    file_like.seek(0)
    reader = PdfReader(file_like)
    pages = []
    for p in reader.pages:
        t = p.extract_text()
        if t:
            pages.append(t)
    return "\n\n".join(pages)

def extract_images_from_pdf(file_like, min_width=300):
    images = []
    try:
        file_like.seek(0)
        reader = PdfReader(file_like)
        for i, page in enumerate(reader.pages):
            if not hasattr(page, "images"):
                continue
            for img in page.images:
                try:
                    pil = PILImage.open(io.BytesIO(img.data))
                    if pil.mode not in ("RGB", "RGBA"):
                        pil = pil.convert("RGB")
                    if pil.size[0] < min_width:
                        continue
                    buf = io.BytesIO()
                    pil.save(buf, format="PNG")
                    buf.seek(0)
                    images.append((f"Page {i+1}", buf))
                except:
                    continue
    except:
        pass
    return images

# =========================================================
# RAG CORE
# =========================================================

def build_vectorstore(uploaded_files):
    docs = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)

    for f in uploaded_files:
        text = extract_text_from_pdf(f)
        chunks = splitter.split_text(text)
        for c in chunks:
            docs.append(
                Document(
                    page_content=c,
                    metadata={"source": f.name}
                )
            )

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )
    return FAISS.from_documents(docs, embeddings)

def rag_retrieve(vectorstore, query, k=15):
    docs = vectorstore.similarity_search(query, k=k)
    return "\n\n".join(d.page_content for d in docs)

# =========================================================
# GROQ ANALYSIS (GAPS + IDEAS)
# =========================================================

def groq_generate(client, prompt, model):
    res = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model=model,
        temperature=0.4,
        max_tokens=2000
    )
    return res.choices[0].message.content

# =========================================================
# GEMINI IEEE PAPER GENERATION (RAG-GROUNDED)
# =========================================================

def generate_ieee_paper(topic, literature_context, references_context):
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY not found in environment")

    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel("gemini-2.5-flash")

    def run_prompt(prompt):
        response = model.generate_content(prompt)
        if not response or not response.text:
            raise RuntimeError("Gemini returned empty response")
        return response.text

    part1 = run_prompt(f"""
You are writing an IEEE research paper.

STRICT RULES:
- Use ONLY the provided context
- DO NOT invent references
- Follow IEEE format

TOPIC:
{topic}

===== LITERATURE CONTEXT START =====
{literature_context[:8000]}
===== LITERATURE CONTEXT END =====

Write:
Title
Abstract (200 words)
I. Introduction
""")

    part2 = run_prompt(f"""
Continue IEEE paper on topic "{topic}"

===== CONTEXT =====
{literature_context[:8000]}

Write:
II. Literature Survey
III. Proposed Methodology
""")

    part3 = run_prompt(f"""
Finish IEEE paper.

===== REFERENCES CONTEXT =====
{references_context[:6000]}

Write:
IV. Expected Results
V. Conclusion
VI. References (IEEE format, ONLY from context)
""")

    return "\n\n".join([part1, part2, part3])


# =========================================================
# DOCX EXPORT
# =========================================================

def create_docx_report(title, content, images=None):
    doc = DocxDocument()
    doc.add_heading(title, 0)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith(("I.", "II.", "III.", "IV.", "V.", "VI.", "Abstract", "Title")):
            doc.add_heading(line, level=2)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if images:
        doc.add_page_break()
        doc.add_heading("Appendix", level=1)
        for lbl, img in images:
            doc.add_paragraph(lbl)
            doc.add_picture(img, width=Inches(5))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# =========================================================
# STREAMLIT APP
# =========================================================

def main():
    st.set_page_config("RAG Research Paper Generator", layout="wide")
    st.title("📚 RAG-Based Automated Literature Review & IEEE Paper Generator")

    if not GROQ_API_KEY:
        st.error("GROQ_API_KEY missing in .env")
        st.stop()

    model_choice = st.sidebar.selectbox(
        "Groq Model",
        ["llama-3.3-70b-versatile", "mixtral-8x7b-32768"]
    )

    uploaded_files = st.file_uploader(
        "Upload IEEE Research Papers (PDF)",
        type=["pdf"],
        accept_multiple_files=True
    )

    if uploaded_files and "vectorstore" not in st.session_state:
        with st.spinner("Building RAG index..."):
            st.session_state.vectorstore = build_vectorstore(uploaded_files)
            st.success("Papers indexed successfully")

    if "vectorstore" not in st.session_state:
        st.info("Upload papers to begin")
        return

    client = Groq(api_key=GROQ_API_KEY)
    tab1, tab2, tab3, tab4, tab5 = st.tabs(
        ["📑 Review", "💡 Gaps & Ideas", "📄 IEEE Paper", "📊 Figures", "🔍 Ask Papers"]
    )

    # ---------------- REVIEW ----------------
    with tab1:
        if st.button("Generate Literature Review"):
            ctx = rag_retrieve(
                st.session_state.vectorstore,
                "Summarize problem statements, methods, datasets, and metrics used across all papers."
            )
            st.session_state.review = groq_generate(
                client,
                f"Generate a unified literature review:\n{ctx}",
                model_choice
            )
        if "review" in st.session_state:
            st.markdown(st.session_state.review)

    # ---------------- GAPS & IDEAS ----------------
    with tab2:
        if st.button("Generate Research Gaps & Novel Ideas"):
            ctx = rag_retrieve(
                st.session_state.vectorstore,
                "Identify limitations, future work, and unexplored research directions."
            )
            st.session_state.gaps = groq_generate(
                client,
                f"""
From the following context:
{ctx}

Provide:
1. Research gaps
2. 3 Novel research ideas with titles and short descriptions
""",
                model_choice
            )
        if "gaps" in st.session_state:
            st.markdown(st.session_state.gaps)

    # ---------------- IEEE PAPER ----------------
    with tab3:
        st.subheader("📄 IEEE Research Paper Generator")

        topic = st.text_input(
            "Select / Paste Research Idea Title",
            placeholder="e.g., Multimodal Stress Detection using Transformer Networks"
        )

        if st.button("🚀 Generate IEEE Paper"):
            if not topic:
                st.warning("Please enter a research topic.")
            elif not GEMINI_API_KEY:
                st.error("GEMINI_API_KEY missing in .env file.")
            else:
                with st.spinner("Gemini is generating the IEEE paper (this may take ~30 seconds)..."):
                    try:
                        literature_ctx = rag_retrieve(
                            st.session_state.vectorstore,
                            "Provide key literature insights relevant to the selected topic."
                        )

                        references_ctx = rag_retrieve(
                            st.session_state.vectorstore,
                            "Extract references including authors, title, venue, and year."
                        )

                        paper = generate_ieee_paper(
                            topic,
                            literature_ctx,
                            references_ctx
                        )   
                        st.session_state.paper = paper
                        st.success("IEEE paper generated successfully!")

                    except Exception as e:
                        st.error(f"IEEE paper generation failed: {e}")

        if "paper" in st.session_state:
            with st.expander("📖 Preview IEEE Paper"):
                st.markdown(st.session_state.paper)

            docx = create_docx_report(
                "IEEE Research Paper",
                st.session_state.paper
            )

            st.download_button(
                "📥 Download IEEE Paper (.docx)",
                docx,
                "IEEE_Paper.docx"
            )

    # ---------------- FIGURES ----------------
    with tab4:
        imgs = []
        for f in uploaded_files:
            imgs.extend(extract_images_from_pdf(f))
        for lbl, img in imgs:
            st.image(img, caption=lbl, width=300)

    # ---------------- ASK PAPERS ----------------
    with tab5:
        q = st.text_input("Ask a question across all papers")
        if q:
            ctx = rag_retrieve(st.session_state.vectorstore, q)
            ans = groq_generate(client, ctx, model_choice)
            st.markdown(ans)

if __name__ == "__main__":
    main()
