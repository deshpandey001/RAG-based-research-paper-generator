import io
import os
import logging

import streamlit as st
from groq import Groq
import google.generativeai as genai
from PIL import Image as PILImage

# ------------------ ENV SETUP ------------------
from dotenv import load_dotenv
load_dotenv()

try:
    GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
except Exception:
    GROQ_API_KEY = os.getenv("GROQ_API_KEY")
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
except Exception:
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

# ------------------ RAG IMPORTS ------------------
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

# ------------------ PDF PARSING ------------------
try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader

# ------------------ DOCX EXPORT ------------------
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------ LOGGING ------------------
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# =========================================================
# PDF HELPERS
# =========================================================

def extract_text_from_pdf(file_like):
    file_like.seek(0)
    reader = PdfReader(file_like)
    pages = []
    for p in reader.pages:
        t = p.extract_text()
        if t:
            pages.append(t)
    return "\n\n".join(pages)


def extract_images_from_pdf(file_like, min_width=300):
    images = []
    try:
        file_like.seek(0)
        reader = PdfReader(file_like)
        for i, page in enumerate(reader.pages):
            if not hasattr(page, "images"):
                continue
            for img in page.images:
                try:
                    pil = PILImage.open(io.BytesIO(img.data))
                    if pil.mode not in ("RGB", "RGBA"):
                        pil = pil.convert("RGB")
                    if pil.size[0] < min_width:
                        continue
                    buf = io.BytesIO()
                    pil.save(buf, format="PNG")
                    buf.seek(0)
                    images.append((f"Page {i+1}", buf))
                except Exception:
                    continue
    except Exception:
        pass
    return images


# =========================================================
# RAG CORE
# =========================================================

def build_vectorstore(uploaded_files):
    docs = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)

    for f in uploaded_files:
        text = extract_text_from_pdf(f)
        chunks = splitter.split_text(text)
        for c in chunks:
            docs.append(
                Document(
                    page_content=c,
                    metadata={"source": f.name},
                )
            )

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )
    return FAISS.from_documents(docs, embeddings)


def rag_retrieve(vectorstore, query, k=15):
    docs = vectorstore.similarity_search(query, k=k)
    return "\n\n".join(d.page_content for d in docs)


# =========================================================
# GROQ ANALYSIS (REVIEW + GAPS + QA)
# =========================================================

def groq_generate(client, prompt, model):
    res = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model=model,
        temperature=0.4,
        max_tokens=2000,
    )
    return res.choices[0].message.content


# =========================================================
# GEMINI WHITE PAPER GENERATION (RAG-GROUNDED)
# =========================================================

def generate_white_paper(topic, domain, literature_context, references_context):
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY not found in environment")

    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel("gemini-2.5-flash")

    def run_prompt(prompt: str) -> str:
        response = model.generate_content(prompt)
        if not response or not getattr(response, "text", None):
            raise RuntimeError("Gemini returned empty response")
        return response.text

    base_lit = literature_context[:8000]
    base_ref = references_context[:6000]

    part1 = run_prompt(
        f"""
You are writing a professional white paper on the use of AI/ML in healthcare.

DOMAIN: {domain}
TOPIC: {topic}

STRICT RULES:
- Use ONLY the provided context
- Be practical, clinically relevant, and persuasive
- Focus on healthcare stakeholders (CMO, CIO, clinicians, hospital admins, payers)
- Highlight impact on patient outcomes, clinical workflows, cost, and risk

===== CONTEXT START =====
{base_lit}
===== CONTEXT END =====

Write:
1. Title
2. Executive Summary (1–2 paragraphs)
3. Clinical / Operational Problem Statement & Challenges
4. Background & Current Landscape of AI/ML in healthcare for this topic
"""
    )

    part2 = run_prompt(
        f"""
Continue the same white paper on AI/ML in healthcare.

DOMAIN: {domain}
TOPIC: {topic}

Use the same context plus the following reference details if useful.

===== LITERATURE CONTEXT =====
{base_lit}

===== REFERENCE / EVIDENCE CONTEXT =====
{base_ref}

Write (focus on healthcare setting):
5. Proposed AI/ML Solution / Approach (algorithms, data sources, workflows)
6. High-level Technical & Clinical Architecture / Methodology
7. Key Benefits and Differentiators (clinical, operational, financial)
8. Key Use Cases / Scenarios (e.g., diagnosis support, triage, risk prediction, resource optimization)
"""
    )

    part3 = run_prompt(
        f"""
Finish the white paper on AI/ML in healthcare.

DOMAIN: {domain}
TOPIC: {topic}

Use ONLY the provided context.

===== LITERATURE CONTEXT =====
{base_lit}

===== REFERENCE / EVIDENCE CONTEXT =====
{base_ref}

Write (keep it grounded in healthcare reality):
9. Implementation Considerations / Adoption Roadmap (data, integration with HIS/EMR, change management)
10. Risks, Limitations, and Mitigations (bias, privacy, security, regulatory, clinical safety)
11. Conclusion and Call to Action for healthcare organizations
12. Optional References (only if they clearly exist in the context; do not invent citations)
"""
    )

    return "\n\n".join([part1, part2, part3])


# =========================================================
# DOCX EXPORT
# =========================================================

def create_docx_report(title, content, images=None):
    doc = DocxDocument()
    doc.add_heading(title, 0)

    heading_starts = (
        "Title",
        "Executive Summary",
        "Clinical / Operational Problem Statement",
        "Problem Statement",
        "Business Challenges",
        "Background",
        "Proposed Solution",
        "High-level Technical & Clinical Architecture",
        "High-level Architecture",
        "Key Benefits",
        "Benefits and Differentiators",
        "Use Cases",
        "Implementation Considerations",
        "Adoption Roadmap",
        "Risks, Limitations, and Mitigations",
        "Risks",
        "Conclusion",
        "References",
    )

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith(heading_starts):
            doc.add_heading(line, level=2)
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if images:
        doc.add_page_break()
        doc.add_heading("Appendix", level=1)
        for lbl, img in images:
            doc.add_paragraph(lbl)
            doc.add_picture(img, width=Inches(5))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# =========================================================
# STREAMLIT APP
# =========================================================

def main():
    st.set_page_config("AI/ML in Healthcare White Paper Generator", layout="wide")
    st.title("📄 RAG-Based AI/ML in Healthcare White Paper Generator")

    if not GROQ_API_KEY:
        st.error("GROQ_API_KEY missing in .env or Streamlit secrets")
        st.stop()

    # Fixed domain for this app: AI/ML in Healthcare
    domain = "AI/ML in Healthcare"
    st.sidebar.markdown("**Domain:** AI/ML in Healthcare")

    model_choice = st.sidebar.selectbox(
        "Groq Model",
        ["llama-3.3-70b-versatile", "mixtral-8x7b-32768"],
    )

    uploaded_files = st.file_uploader(
        "Upload Background Material (clinical papers, healthcare AI reports, policy docs as PDF)",
        type=["pdf"],
        accept_multiple_files=True,
    )

    if uploaded_files and "vectorstore_wp" not in st.session_state:
        with st.spinner("Building RAG index for white paper..."):
            st.session_state.vectorstore_wp = build_vectorstore(uploaded_files)
            st.success("Documents indexed successfully for white paper generation")

    if "vectorstore_wp" not in st.session_state:
        st.info("Upload documents to begin generating a white paper")
        return

    client = Groq(api_key=GROQ_API_KEY)

    tab1, tab2, tab3, tab4, tab5 = st.tabs(
        [
            "📑 Literature Review",
            "💡 Gaps & Opportunities",
            "📄 White Paper",
            "📊 Figures",
            "🔍 Ask Documents",
        ]
    )

    # ---------------- REVIEW ----------------
    with tab1:
        if st.button("Generate Literature Review", key="wp_review_btn"):
            ctx = rag_retrieve(
                st.session_state.vectorstore_wp,
                "Summarize problem statements, approaches, datasets, and key findings across all documents.",
            )
            st.session_state.wp_review = groq_generate(
                client,
                f"Generate a cohesive, domain-aware literature review for domain '{domain}' based on the following context:\n{ctx}",
                model_choice,
            )
        if "wp_review" in st.session_state:
            st.markdown(st.session_state.wp_review)

    # ---------------- GAPS & OPPORTUNITIES ----------------
    with tab2:
        if st.button("Generate Gaps & Opportunities", key="wp_gaps_btn"):
            ctx = rag_retrieve(
                st.session_state.vectorstore_wp,
                "Identify limitations, open challenges, and future opportunities mentioned across the documents.",
            )
            st.session_state.wp_gaps = groq_generate(
                client,
                f"""
You are analysing documents for the domain: {domain}.

From the following context:
{ctx}

Provide:
1. Key limitations and pain points in current approaches
2. Open challenges and unmet needs
3. 3–5 concrete opportunity areas suitable for a new solution or product
""",
                model_choice,
            )
        if "wp_gaps" in st.session_state:
            st.markdown(st.session_state.wp_gaps)

    # ---------------- WHITE PAPER GENERATION ----------------
    with tab3:
        st.subheader("📄 Domain-Focused White Paper Generator")

        topic = st.text_input(
            "Enter White Paper Topic / Solution Title",
            placeholder="e.g., Zero-Trust Architecture for Enterprise Cloud Security",
        )

        if st.button("🚀 Generate White Paper", key="wp_generate_btn"):
            if not topic:
                st.warning("Please enter a white paper topic.")
            elif not GEMINI_API_KEY:
                st.error("GEMINI_API_KEY missing in .env or Streamlit secrets.")
            else:
                with st.spinner("Gemini is generating the white paper (this may take ~30 seconds)..."):
                    try:
                        literature_ctx = rag_retrieve(
                            st.session_state.vectorstore_wp,
                            f"Provide key domain-specific insights, trends, and approaches relevant to a white paper on: {topic} in {domain}.",
                        )

                        references_ctx = rag_retrieve(
                            st.session_state.vectorstore_wp,
                            "Extract concrete facts, statistics, case studies, and any explicit references (authors, titles, venues, years).",
                        )

                        white_paper = generate_white_paper(
                            topic,
                            domain,
                            literature_ctx,
                            references_ctx,
                        )
                        st.session_state.white_paper = white_paper
                        st.success("White paper generated successfully!")

                    except Exception as e:
                        st.error(f"White paper generation failed: {e}")

        if "white_paper" in st.session_state:
            with st.expander("📖 Preview White Paper"):
                st.markdown(st.session_state.white_paper)

            docx = create_docx_report(
                "White Paper",
                st.session_state.white_paper,
            )

            st.download_button(
                "📥 Download White Paper (.docx)",
                docx,
                "White_Paper.docx",
            )

    # ---------------- FIGURES ----------------
    with tab4:
        imgs = []
        if uploaded_files:
            for f in uploaded_files:
                imgs.extend(extract_images_from_pdf(f))
        if imgs:
            for lbl, img in imgs:
                st.image(img, caption=lbl, width=300)
        else:
            st.info("No suitable figures detected in uploaded PDFs.")

    # ---------------- ASK DOCUMENTS ----------------
    with tab5:
        q = st.text_input("Ask a question across all uploaded documents")
        if q:
            ctx = rag_retrieve(st.session_state.vectorstore_wp, q)
            ans = groq_generate(
                client,
                f"You are answering a domain-specific question for domain '{domain}'.\n\nContext from documents:\n{ctx}\n\nQuestion: {q}",
                model_choice,
            )
            st.markdown(ans)


if __name__ == "__main__":
    main()
